import os
import threading
import asyncio
import customtkinter as ctk
from docx import Document
import google.generativeai as genai
from datetime import datetime
from typing import Dict, Tuple, Optional
import traceback

# Define paper sections with their prompts
sections = [
    ("Abstract", 
     "Write a detailed and structured abstract for a research paper about {topic}. Summarize the background, objectives, methods, key results, and conclusions in a concise academic style, including the significance of the findings."),
    ("Introduction", 
     "Write a comprehensive introduction for a research paper on {topic}. Include the study's significance, relevant background, existing research gaps, and how this study addresses them. Use an engaging and academic tone."),
    ("Methodology", 
     "Provide an in-depth methodology section for a research paper about {topic}. Include participant demographics, data collection instruments, procedures, and data analysis methods. Mention ethical considerations if applicable."),
    ("Results", 
     "Elaborate on the research findings for {topic}. Include detailed statistical results, tables if applicable, subgroup analyses, and comparisons to existing research. Ensure clear explanations of statistical terms."),
    ("Discussion", 
     "Write a critical discussion of the results obtained in the study about {topic}. Analyze implications, strengths, limitations, compare with related literature, and suggest future research directions."),
    ("Conclusion", 
     "Summarize the key findings and their significance for the study on {topic}. Discuss potential applications of the research, limitations, and recommendations for future studies.")
]

def build_context(section_name: str, previous_sections: Dict[str, str], citations: Dict[str, Dict[str, str]] = None) -> str:
    """Build context from previous sections for coherent paper generation with citations"""
    # Format citations if they exist
    citation_text = ""
    if citations and section_name in citations:
        citation_text = ", ".join(citations[section_name].values())
    
    contexts = {
        'Abstract': "",  # Abstract needs no previous context
        'Introduction': """Previously in the Abstract, we established the research focus on {topic} 
            using {methodology} which yielded {results}. Include relevant citations: {citations}.""",
        'Methodology': """Based on the research objectives outlined in the Introduction regarding {topic},
            detail the following methodological approach: {methodology}. Reference similar methodologies: {citations}.""",
        'Results': """Following the {methodology} described in the Methodology section,
            present these findings: {results}. Compare with similar studies: {citations}.""",
        'Discussion': """Considering the results showing {results} obtained through {methodology},
            interpret these findings in the context of {topic}. Include comparative analysis: {citations}.""",
        'Conclusion': """Given the findings {results} and their discussion in relation to {topic},
            provide concluding thoughts. Reference key literature: {citations}."""
    }
    
    base_context = contexts.get(section_name, "")
    if not base_context:
        return ""
        
    return base_context.format(
        topic=previous_sections.get('topic', ''),
        methodology=previous_sections.get('methodology', ''),
        results=previous_sections.get('results', ''),
        citations=citation_text
    )

class ResearchPaperGeneratorApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # Initialize variables
        self.citations: Dict[str, Dict[str, str]] = {}
        self.current_task = None
        self.closing = False
        self.inputs = {}
        self.progress_bars = {}
        self.status_labels = {}
        self.section_texts = {}
        self.save_as_markdown = False
        
        # Configure window
        self.title("Research Paper Generator")
        self.geometry("800x600")
        
        # Setup UI
        self.setup_ui()

    def setup_ui(self):
        # Configure grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=3)

        # Create main frame
        main_frame = ctk.CTkFrame(self)
        main_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        main_frame.grid_columnconfigure(1, weight=1)

        # Input fields
        row = 0
        # API Key
        ctk.CTkLabel(main_frame, text="API Key:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        self.inputs['api_key'] = ctk.CTkEntry(main_frame, show="*", width=400)
        self.inputs['api_key'].grid(row=row, column=1, padx=5, pady=5, sticky="ew")
        
        row += 1
        # Topic
        ctk.CTkLabel(main_frame, text="Research Topic:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        self.inputs['topic'] = ctk.CTkEntry(main_frame, width=400)
        self.inputs['topic'].grid(row=row, column=1, padx=5, pady=5, sticky="ew")
        
        row += 1
        # Methodology
        ctk.CTkLabel(main_frame, text="Methodology:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        self.inputs['methodology'] = ctk.CTkEntry(main_frame, width=400)
        self.inputs['methodology'].grid(row=row, column=1, padx=5, pady=5, sticky="ew")
        
        row += 1
        # Results
        ctk.CTkLabel(main_frame, text="Key Results:").grid(row=row, column=0, padx=5, pady=5, sticky="e")
        self.inputs['results'] = ctk.CTkEntry(main_frame, width=400)
        self.inputs['results'].grid(row=row, column=1, padx=5, pady=5, sticky="ew")

        row += 1
        # Save as Markdown option
        self.save_as_markdown_var = ctk.BooleanVar()
        ctk.CTkCheckBox(main_frame, text="Save as Markdown", variable=self.save_as_markdown_var).grid(row=row, column=0, columnspan=2, pady=5)

        row += 1
        # Generate button
        self.generate_button = ctk.CTkButton(
            main_frame, 
            text="Generate Paper", 
            command=self.start_generation
        )
        self.generate_button.grid(row=row, column=0, columnspan=2, pady=20)

        row += 1
        # Status label
        self.status_label = ctk.CTkLabel(main_frame, text="")
        self.status_label.grid(row=row, column=0, columnspan=2, pady=5)

        # Progress section
        row += 1
        progress_frame = ctk.CTkFrame(main_frame)
        progress_frame.grid(row=row, column=0, columnspan=2, padx=10, pady=10, sticky="ew")
        progress_frame.grid_columnconfigure(1, weight=1)

        # Create progress bars and labels for each section
        for i, (section_name, _) in enumerate(sections):
            ctk.CTkLabel(progress_frame, text=f"{section_name}:").grid(row=i, column=0, padx=5, pady=2, sticky="e")
            
            progress_bar = ctk.CTkProgressBar(progress_frame)
            progress_bar.grid(row=i, column=1, padx=5, pady=2, sticky="ew")
            progress_bar.set(0)
            self.progress_bars[section_name] = progress_bar
            
            status_label = ctk.CTkLabel(progress_frame, text="Pending")
            status_label.grid(row=i, column=2, padx=5, pady=2)
            self.status_labels[section_name] = status_label

    def validate_inputs(self) -> Tuple[bool, str]:
        """Validate all input fields"""
        required_fields = ['api_key', 'topic', 'methodology', 'results']
        for field in required_fields:
            if not self.inputs[field].get().strip():
                return False, f"{field.title()} is required"
        return True, ""

    def update_status_label(self, message: str, is_error: bool = False):
        """Update the main status label"""
        color = "red" if is_error else "white"
        self.status_label.configure(text=message, text_color=color)

    def update_progress(self, section: str, status: str, progress: float):
        """Update progress bar and status for a section"""
        if section in self.progress_bars:
            self.progress_bars[section].set(progress)
            self.status_labels[section].configure(text=status)

    def toggle_inputs(self, enabled: bool):
        """Enable or disable input fields and generate button"""
        state = "normal" if enabled else "disabled"
        for input_field in self.inputs.values():
            input_field.configure(state=state)
        self.generate_button.configure(state=state)

    def start_generation(self):
        """Start the paper generation process"""
        if self.current_task:
            return

        self.toggle_inputs(False)
        self.update_status_label("Generating paper...")

        # Reset progress bars and storage
        self.citations.clear()
        self.section_texts.clear()
        for section_name in self.progress_bars:
            self.update_progress(section_name, "Pending", 0)

        # Start generation in a separate thread
        self.current_task = threading.Thread(target=self.run_async_generation)
        self.current_task.start()

    def run_async_generation(self):
        """Run the async generation process in a separate thread"""
        asyncio.run(self.generate_paper())

    def save_document(self, section_texts: Dict[str, str]):
        """Save the generated paper to a Word document"""
        doc = Document()
        for section_name, text in section_texts.items():
            doc.add_heading(section_name, level=1)
            doc.add_paragraph(text)
            doc.add_paragraph()  # Add spacing between sections

        # Save with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_word = f"research_paper_{timestamp}.docx"
        doc.save(filename_word)

        # Save to Markdown file if option is selected
        if self.save_as_markdown_var.get():
            self.save_as_markdown = True
            filename_md = f"research_paper_{timestamp}.md"
            formatted_text = format_to_markdown(section_texts)
            with open(filename_md, "w", encoding="utf-8") as md_file:
                md_file.write(formatted_text)

        self.update_status_label(f"Paper saved as {filename_word} and {filename_md if self.save_as_markdown else ''}")

    async def generate_paper(self):
        """Generate the research paper section by section"""
        try:
            # Validate inputs
            valid, message = self.validate_inputs()
            if not valid:
                self.update_status_label(message, is_error=True)
                return

            # Configure API
            api_key = self.inputs['api_key'].get()
            os.environ["GEMINI_API_KEY"] = api_key
            genai.configure(api_key=api_key)

            # Generation configuration
            generation_config = {
                "temperature": 0.001,
                "top_p": 1,
                "top_k": 40,
                "max_output_tokens": 8192,
            }

            # Configure model
            model = genai.GenerativeModel(
                model_name="gemini-1.5-flash",
                generation_config=generation_config
            )

            # Store input values for context building
            input_values = {
                'topic': self.inputs['topic'].get(),
                'methodology': self.inputs['methodology'].get(),
                'results': self.inputs['results'].get()
            }

            for section_name, prompt_template in sections:
                if self.closing:
                    break

                self.update_progress(section_name, "Generating...", 0.5)
                
                try:
                    # Build context with citations
                    context = build_context(section_name, input_values, self.citations)
                    
                    # Create prompt
                    prompt = f"""{context}
                    
                    Requirements:
                    1. Use academic writing style
                    2. Include proper citations
                    3. Maintain consistent formatting
                    4. Ensure factual accuracy
                    5. Cross-reference with existing literature
                    
                    {prompt_template.format(**input_values)}
                    """

                    # Generate content
                    response = model.generate_content(prompt)
                    section_text = response.text
                    
                    # Extract and store citations
                    self.extract_citations(section_text, section_name)
                    
                    self.section_texts[section_name] = section_text
                    self.update_progress(section_name, "Complete", 1.0)
                    
                    # Short delay between sections
                    await asyncio.sleep(1)
                
                except Exception as e:
                    self.update_progress(section_name, "Failed", 0)
                    raise Exception(f"Error generating {section_name}: {str(e)}")

            if not self.closing:
                self.save_document(self.section_texts)
                self.update_status_label("Research paper generated successfully!")

        except Exception as e:
            self.update_status_label(f"Error: {str(e)}", is_error=True)
            print(traceback.format_exc())  # Print full traceback for debugging
        
        finally:
            self.toggle_inputs(True)
            self.current_task = None

    def extract_citations(self, text: str, section_name: str):
        """Extract and store citations from generated text"""
        import re
        citations = re.findall(r'\(([^)]+)\)', text)
        if citations:
            self.citations[section_name] = {f"citation_{i}": citation for i, citation in enumerate(citations)}

def format_to_markdown(section_texts):
    markdown_sections = []
    for section, content in section_texts.items():
        markdown_sections.append(f"## {section}\n\n{content.strip()}\n")
    return "\n---\n".join(markdown_sections)

if __name__ == "__main__":
    app = ResearchPaperGeneratorApp()
    app.mainloop()